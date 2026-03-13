from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import random
from functools import lru_cache

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf"}
_CANDIDATE_RE = re.compile(r"^(cand_\d+)_(cv|cover)$", re.IGNORECASE)
_JOB_RE = re.compile(r"^jd_\d+$", re.IGNORECASE)


@dataclass(frozen=True)
class Document:
    doc_id: str
    text: str
    path: Path


@dataclass(frozen=True)
class CandidateDoc:
    candidate_id: str
    kind: str
    text: str
    path: Path


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_docx(path: Path) -> str:
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs)


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        if txt.strip():
            parts.append(txt)
    return "\n".join(parts)


def read_document(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".txt":
        return read_txt(path)
    if ext == ".docx":
        return read_docx(path)
    if ext == ".pdf":
        return read_pdf(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


@lru_cache(maxsize=16)
def _cached_supported_files(folder_str: str) -> tuple[str, ...]:
    folder = Path(folder_str)
    if not folder.exists():
        return tuple()
    return tuple(
        sorted(
            str(p)
            for p in folder.iterdir()
            if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
        )
    )


def list_supported_files(folder: Path) -> list[Path]:
    return [Path(p) for p in _cached_supported_files(str(folder.resolve()))]


@lru_cache(maxsize=5000)
def _read_document_cached(path_str: str, mtime_ns: int, size: int) -> str:
    del mtime_ns, size
    return read_document(Path(path_str))


def read_document_cached(path: Path) -> str:
    stat = path.stat()
    return _read_document_cached(str(path.resolve()), int(stat.st_mtime_ns), int(stat.st_size))


def _sample_paths(paths: list[Path], max_items: int | None, seed: int) -> list[Path]:
    if max_items is None or max_items <= 0 or len(paths) <= max_items:
        return paths
    rng = random.Random(seed)
    sampled = rng.sample(paths, max_items)
    return sorted(sampled)


def count_dataset_items(folder: Path, pattern: str, unique_group: int | None = None) -> int:
    if not folder.exists():
        return 0
    rx = re.compile(pattern, flags=re.IGNORECASE)
    if unique_group is None:
        count = 0
        for p in folder.iterdir():
            if not p.is_file() or p.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            if rx.match(p.stem):
                count += 1
        return count

    unique_values: set[str] = set()
    for p in folder.iterdir():
        if not p.is_file() or p.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        m = rx.match(p.stem)
        if not m:
            continue
        unique_values.add(m.group(unique_group).lower())
    return len(unique_values)


def load_jobs(folder: Path, max_items: int | None = None, seed: int = 42) -> list[Document]:
    job_paths = [p for p in list_supported_files(folder) if _JOB_RE.match(p.stem)]
    sampled_paths = _sample_paths(job_paths, max_items=max_items, seed=seed)
    docs: list[Document] = []
    for p in sampled_paths:
        text = read_document_cached(p).strip()
        if text:
            docs.append(Document(doc_id=p.stem, text=text, path=p))
    return docs


def load_candidate_docs(folder: Path, max_items: int | None = None, seed: int = 42) -> list[CandidateDoc]:
    paths = list_supported_files(folder)
    candidates_to_keep: set[str] | None = None
    if max_items is not None and max_items > 0:
        candidate_ids: list[str] = []
        for p in paths:
            m = _CANDIDATE_RE.match(p.stem)
            if not m:
                continue
            if m.group(2).lower() != "cv":
                continue
            candidate_ids.append(m.group(1).lower())
        candidate_ids = sorted(set(candidate_ids))
        if len(candidate_ids) > max_items:
            rng = random.Random(seed)
            candidates_to_keep = set(rng.sample(candidate_ids, max_items))
        else:
            candidates_to_keep = set(candidate_ids)

    docs: list[CandidateDoc] = []
    for p in paths:
        m = _CANDIDATE_RE.match(p.stem)
        if not m:
            continue
        candidate_id = m.group(1).lower()
        if candidates_to_keep is not None and candidate_id not in candidates_to_keep:
            continue
        kind = m.group(2).lower()
        text = read_document_cached(p).strip()
        if text:
            docs.append(CandidateDoc(candidate_id=candidate_id, kind=kind, text=text, path=p))
    return docs
