from __future__ import annotations

from pathlib import Path
from typing import Iterable
import re

from docx import Document as DocxDocument


def _candidate_label(candidate_id: str) -> str:
    # Upload-mode generated IDs are rendered as Candidate X in reports.
    m = re.match(r"^upload_cand_(\d+)$", candidate_id)
    if not m:
        return candidate_id
    return f"Candidate {int(m.group(1))}"


def write_report_docx(report_path: Path, mode: str, top_k: int, rows: list[dict]) -> None:
    doc = DocxDocument()
    doc.add_heading("Candidate Matching Report", level=1)
    doc.add_paragraph(f"Mode: {mode}")
    doc.add_paragraph(f"Top-K per job: {top_k}")

    by_job: dict[str, list[dict]] = {}
    for row in rows:
        by_job.setdefault(str(row["job_id"]), []).append(row)

    for job_id in sorted(by_job.keys()):
        job_title = str(by_job[job_id][0].get("job_title", "")).strip()
        heading = f"Job: {job_title} ({job_id})" if job_title else f"Job: {job_id}"
        doc.add_heading(heading, level=2)
        for r in sorted(by_job[job_id], key=lambda x: int(x["rank"])):
            candidate_label = _candidate_label(str(r["candidate_id"]))
            paragraph = doc.add_paragraph(
                f"#{r['rank']}  Candidate: {candidate_label}  Similarity: {float(r['similarity']):.4f}"
            )
            if "explanation" in r and str(r["explanation"]).strip():
                paragraph.add_run(f"\nExplanation: {str(r['explanation']).strip()}")

    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def write_report_txt(report_path: Path, mode: str, top_k: int, rows: Iterable[dict]) -> None:
    lines = [f"Candidate Matching Report", f"Mode: {mode}", f"Top-K per job: {top_k}", ""]
    by_job: dict[str, list[dict]] = {}
    for row in rows:
        by_job.setdefault(str(row["job_id"]), []).append(row)

    for job_id in sorted(by_job.keys()):
        job_title = str(by_job[job_id][0].get("job_title", "")).strip()
        heading = f"Job: {job_title} ({job_id})" if job_title else f"Job: {job_id}"
        lines.append(heading)
        for r in sorted(by_job[job_id], key=lambda x: int(x["rank"])):
            candidate_label = _candidate_label(str(r["candidate_id"]))
            lines.append(f"  #{r['rank']}  {candidate_label}  {float(r['similarity']):.4f}")
            if "explanation" in r and str(r["explanation"]).strip():
                lines.append(f"    Explanation: {str(r['explanation']).strip()}")
        lines.append("")

    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
